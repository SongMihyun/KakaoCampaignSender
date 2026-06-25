from __future__ import annotations

import os
import tempfile
from datetime import datetime
from typing import Iterable, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ContactRow = Tuple[str, str, str, str, str, str, str, str, str, str]
HEADERS = [
    "카카오톡 검색명",
    "고객명",
    "호칭",
    "직책",
    "소속/대리점",
    "지사",
    "연락처",
    "상태",
    "태그",
    "메모",
]


def _apply_sheet_style(ws) -> None:
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="2F5597")
    header_align = Alignment(horizontal="center", vertical="center")

    ws.append(HEADERS)
    ws.freeze_panes = "A2"

    for col in range(1, len(HEADERS) + 1):
        cell = ws.cell(row=1, column=col)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    ws.row_dimensions[1].height = 20

    widths = [22, 14, 10, 12, 18, 14, 16, 12, 16, 26]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w


def _atomic_save_workbook(wb: Workbook, path: str) -> None:
    path = os.path.abspath(path)
    folder = os.path.dirname(path) or "."
    os.makedirs(folder, exist_ok=True)

    fd, tmp_path = tempfile.mkstemp(prefix=".__tmp_", suffix=".xlsx", dir=folder)
    os.close(fd)

    try:
        wb.save(tmp_path)
        os.replace(tmp_path, path)
    except Exception:
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass
        raise


def _atomic_save_docx(doc: Document, path: str) -> None:
    path = os.path.abspath(path)
    folder = os.path.dirname(path) or "."
    os.makedirs(folder, exist_ok=True)

    fd, tmp_path = tempfile.mkstemp(prefix=".__tmp_", suffix=".docx", dir=folder)
    os.close(fd)

    try:
        doc.save(tmp_path)
        os.replace(tmp_path, path)
    except Exception:
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass
        raise


def _atomic_write_text(path: str, text: str, encoding: str = "utf-8-sig") -> None:
    path = os.path.abspath(path)
    folder = os.path.dirname(path) or "."
    os.makedirs(folder, exist_ok=True)

    fd, tmp_path = tempfile.mkstemp(prefix=".__tmp_", suffix=os.path.splitext(path)[1] or ".txt", dir=folder)
    os.close(fd)

    try:
        with open(tmp_path, "w", encoding=encoding, newline="") as f:
            f.write(text)
        os.replace(tmp_path, path)
    except Exception:
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass
        raise


def _sample_rows() -> list[ContactRow]:
    return [
        (
            "홍길동",
            "홍길동",
            "고객님",
            "대표",
            "강남대리점",
            "서울",
            "01011112222",
            "상담중",
            "자동차보험",
            "카카오톡 검색명은 실제 채팅방 검색값입니다.",
        ),
        ("김하늘 매니저님", "김하늘", "고객님", "매니저", "", "", "", "", "", ""),
    ]


def create_template_xlsx(path: str) -> None:
    wb = Workbook()

    ws = wb.active
    ws.title = "대상자"
    _apply_sheet_style(ws)

    for row in _sample_rows():
        ws.append(list(row))

    ws2 = wb.create_sheet("안내")
    ws2["A1"] = "입력 규칙"
    ws2["A1"].font = Font(bold=True)
    ws2["A3"] = "1) 카카오톡 검색명은 실제 카카오톡에서 찾을 이름/채팅방명입니다."
    ws2["A4"] = "2) 고객명은 메시지 개인화에 사용할 실제 고객명입니다."
    ws2["A5"] = "3) 예전 파일처럼 '이름' 컬럼만 있어도 검색명과 고객명에 같이 반영됩니다."
    ws2["A6"] = f"4) 생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

    _atomic_save_workbook(wb, path)


def create_template_docx(path: str) -> None:
    doc = Document()

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, h in enumerate(HEADERS):
        hdr[i].text = h

    for row in _sample_rows():
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v or "")

    _atomic_save_docx(doc, path)


def create_template_txt(path: str) -> None:
    lines = ["\t".join(HEADERS)]
    for row in _sample_rows():
        lines.append("\t".join([str(v or "") for v in row]))
    _atomic_write_text(path, "\n".join(lines))


def export_contacts_xlsx(path: str, rows: Iterable[ContactRow]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "대상자"
    _apply_sheet_style(ws)

    for r in rows:
        ws.append(list(r))

    _atomic_save_workbook(wb, path)


def export_contacts_docx(path: str, rows: Iterable[ContactRow]) -> None:
    doc = Document()
    title = doc.add_heading("대상자 내보내기", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    info = doc.add_paragraph(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    if info.runs:
        info.runs[0].font.size = Pt(10)

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0].cells
    for i, h in enumerate(HEADERS):
        hdr[i].text = h

    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v or "")

    _atomic_save_docx(doc, path)


def export_contacts_txt(path: str, rows: Iterable[ContactRow]) -> None:
    lines = ["\t".join(HEADERS)]
    for row in rows:
        lines.append("\t".join([str(v or "") for v in row]))
    _atomic_write_text(path, "\n".join(lines))
