from __future__ import annotations

import csv
import io
import os
import tempfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as _DocumentClass
from docx.table import Table
from docx.text.paragraph import Paragraph

try:
    from backend.integrations.excel.workbook_editor_io import SheetGrid, WorkbookGrid
except Exception:  # pragma: no cover
    from .workbook_editor_io import SheetGrid, WorkbookGrid

SUPPORTED_WORD_EDITOR_EXTS = {".docx"}
SUPPORTED_TEXT_EDITOR_EXTS = {".txt", ".csv", ".tsv"}
_TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "cp949", "euc-kr")


def is_supported_word_editor_file(path: str) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_WORD_EDITOR_EXTS


def is_supported_text_editor_file(path: str) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_TEXT_EDITOR_EXTS


def load_word_grid(path: str) -> WorkbookGrid:
    if not is_supported_word_editor_file(path):
        raise ValueError("지원 확장자: .docx")

    doc = Document(path)
    rows: list[list[str]] = []
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = _clean_text(block.text)
            if text:
                rows.append([text])
        elif isinstance(block, Table):
            for row in block.rows:
                values = [_clean_text(cell.text) for cell in row.cells]
                while values and values[-1] == "":
                    values.pop()
                rows.append(values if values else [""])

    if not rows:
        rows = [[""]]

    sheet = SheetGrid(name="워드", rows=rows)
    sheet.ensure_rectangular()
    return WorkbookGrid(source_path=path, sheets=[sheet])


def load_text_grid(path: str) -> WorkbookGrid:
    if not is_supported_text_editor_file(path):
        raise ValueError("지원 확장자: .txt, .csv, .tsv")

    text = _read_text_file(path)
    ext = Path(path).suffix.lower()
    rows = _rows_from_text_blob(text, ext)
    if not rows:
        rows = [[""]]

    sheet = SheetGrid(name="메모장", rows=rows)
    sheet.ensure_rectangular()
    return WorkbookGrid(source_path=path, sheets=[sheet])


def save_workbook_grid_to_docx(grid: WorkbookGrid, path: str) -> None:
    doc = Document()

    for s_idx, sheet in enumerate(grid.sheets):
        sheet.ensure_rectangular()
        if s_idx > 0:
            doc.add_page_break()
        if len(grid.sheets) > 1:
            doc.add_heading(sheet.name or f"Sheet{s_idx + 1}", level=2)

        rows = sheet.rows or [[""]]
        row_count = max(1, len(rows))
        col_count = max(1, max((len(r) for r in rows), default=1))
        table = doc.add_table(rows=row_count, cols=col_count)
        table.style = "Table Grid"
        for r_idx in range(row_count):
            row = rows[r_idx] if r_idx < len(rows) else []
            for c_idx in range(col_count):
                value = row[c_idx] if c_idx < len(row) else ""
                table.cell(r_idx, c_idx).text = "" if value is None else str(value)

    _save_document_atomic(doc, path)


def save_workbook_grid_to_text(grid: WorkbookGrid, path: str) -> None:
    ext = Path(path).suffix.lower()
    delimiter = "," if ext == ".csv" else "\t"
    sheet = grid.sheets[0] if grid.sheets else SheetGrid(name="메모장", rows=[[""]])
    sheet.ensure_rectangular()

    buf = io.StringIO()
    writer = csv.writer(buf, delimiter=delimiter, lineterminator="\n")
    for row in sheet.rows:
        trimmed = list(row)
        while trimmed and trimmed[-1] == "":
            trimmed.pop()
        writer.writerow(trimmed)

    _atomic_write_text(path, buf.getvalue())


def suggest_word_save_path(source_path: str) -> str:
    src = Path(source_path)
    if src.suffix.lower() == ".docx":
        return str(src)
    return str(src.with_name(f"{src.stem}_edited.docx"))


def suggest_text_save_path(source_path: str) -> str:
    src = Path(source_path)
    if src.suffix.lower() in SUPPORTED_TEXT_EDITOR_EXTS:
        return str(src)
    return str(src.with_name(f"{src.stem}_edited.txt"))


def _iter_block_items(parent: _DocumentClass):
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def _rows_from_text_blob(text: str, ext: str) -> list[list[str]]:
    lines = [ln.rstrip("\r") for ln in (text or "").splitlines()]
    lines = [ln for ln in lines if ln.strip() != ""]
    if not lines:
        return []

    delimiter = _detect_delimiter(lines, ext)
    rows: list[list[str]] = []
    if delimiter is None:
        for line in lines:
            rows.append([_clean_text(line)])
    else:
        for line in lines:
            parsed = next(csv.reader([line], delimiter=delimiter))
            values = [_clean_text(v) for v in parsed]
            while values and values[-1] == "":
                values.pop()
            rows.append(values if values else [""])
    return rows


def _detect_delimiter(lines: Iterable[str], ext: str) -> str | None:
    if ext == ".csv":
        return ","
    if ext == ".tsv":
        return "\t"
    sample = "\n".join(list(lines)[:10])
    for delim in ("\t", ",", ";", "|"):
        if delim in sample:
            return delim
    return None


def _read_text_file(path: str) -> str:
    last_error: Exception | None = None
    for enc in _TEXT_ENCODINGS:
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError as e:
            last_error = e
            continue
    if last_error:
        raise last_error
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _clean_text(value: str) -> str:
    return str(value or "").replace("\r", "").strip()


def _save_document_atomic(doc: Document, path: str) -> None:
    path = os.path.abspath(path)
    folder = os.path.dirname(path) or "."
    os.makedirs(folder, exist_ok=True)
    fd, tmp_path = tempfile.mkstemp(prefix=".__tmp_", suffix=".docx", dir=folder)
    os.close(fd)
    try:
        doc.save(tmp_path)
        os.replace(tmp_path, path)
    except Exception:
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass
        raise


def _atomic_write_text(path: str, text: str) -> None:
    path = os.path.abspath(path)
    folder = os.path.dirname(path) or "."
    os.makedirs(folder, exist_ok=True)
    fd, tmp_path = tempfile.mkstemp(prefix=".__tmp_", suffix=Path(path).suffix or ".txt", dir=folder)
    os.close(fd)
    try:
        with open(tmp_path, "w", encoding="utf-8-sig", newline="") as f:
            f.write(text)
        os.replace(tmp_path, path)
    except Exception:
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass
        raise
